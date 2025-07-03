from docx import Document


# Chemins des fichiers
docx_path = "Template_mod.docx"
txt_path = 'notes.txt'
output_docx_path = './result.docx'
reference = "reference.xlsx"


# Charger le document Word
doc = Document(docx_path)
new_doc = Document()


# Lire les associations identifiant-valeur depuis le fichier texte
associations = {}
with open(txt_path, 'r', encoding='utf-8') as file:
    for line in file:
        if '::' in line:
            identifiant, valeur = line.strip().split('::', 1)
            associations[identifiant] = valeur

# Parcourir chaque paragraphe du document
for paragraph in doc.paragraphs:
    for identifiant, valeur in associations.items():
        str = "{1" + identifiant + "}"
        if str in paragraph.text:
            paragraph.text = paragraph.text.replace(str, valeur)
    new_doc.add_paragraph(paragraph.text)

# Parcourir chaque section du document pour les en-têtes
    for section in doc.sections:
        for header in section.header.paragraphs:
            for identifiant, valeur in associations.items():
                str = "{1" + identifiant + "}"
                if str in paragraph.text:
                    paragraph.text = paragraph.text.replace(str, valeur)
            new_doc.add_paragraph(paragraph.text)


# Parcourir chaque tableau du document
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for identifiant, valeur in associations.items():
                if f"{{{{{identifiant}}}}}" in cell.text:
                    cell.text = cell.text.replace(f"{{{{{identifiant}}}}}", valeur)


# Sauvegarder le document modifié
new_doc.save(output_docx_path)
