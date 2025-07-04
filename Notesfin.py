from docx import Document
import docs

# Chemins des fichiers
docx_path = docs.template_path
txt_path = docs.txt_path
output_docx_path = docs.sortie_path

# Charger le document Word
doc = Document(docx_path)

# Lire les associations identifiant-valeur depuis le fichier texte
associations = {}
with open(txt_path, 'r', encoding='utf-8') as file:
    for line in file:
        if '::' in line:
            identifiant, valeur = line.strip().split('::', 1)
            associations[identifiant] = valeur

# Fonction pour remplacer les identifiants dans le texte
def replace_identifiers(text):
    for identifiant, valeur in associations.items():
        placeholder = "[" + identifiant + "]"
        if placeholder in text:
            text = text.replace(placeholder, valeur)
    return text

# Parcourir chaque paragraphe du document
for paragraph in doc.paragraphs:
    paragraph.text = replace_identifiers(paragraph.text)

# Parcourir chaque section du document pour les en-têtes
for section in doc.sections:
    for header_paragraph in section.header.paragraphs:
        header_paragraph.text = replace_identifiers(header_paragraph.text)

# Parcourir chaque tableau du document
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.text = replace_identifiers(paragraph.text)

# Sauvegarder le document modifié
doc.save(output_docx_path)
