import pandas as pd 
from docx import Document
from generateur_proposition import proposition
from remplissage_axes import remplissage
import docs
from docx.oxml import OxmlElement
from Notes_word import Word_note

DEVIS = docs.devis_3_path
df_devis = pd.read_excel(DEVIS)
REFERENCE = docs.reference
df_reference = pd.read_excel(REFERENCE)
word = Document(docs.template_path)
Contexte_Reference = proposition() #contexte et ref contexte:texte, reference:liste de texte
Axes_remplis = remplissage() # Titre des axes et sous axes titres:liste de sous axes
test_notes = docs.test_notes_path


def modifie_doc(doc, remplace, remplacant):
    for paragraphe in doc.paragraphs:
        if f"[{remplace}]" in paragraphe.text:
            paragraphe.text = paragraphe.text.replace(remplace, remplacant)

# mets des bullet points 
def modifie_doc_bullet(doc, replace, list):
    for para in doc.paragraphs:
        if replace in para.text:
            # Trouver la position dans le parent XML
            parent = para._element.getparent()
            index = parent.index(para._element)

            # Supprimer le paragraphe contenant le texte à remplacer
            parent.remove(para._element)

            # Insérer les bullet points à la même position
            for bullet in reversed(list):
                new_para = OxmlElement("w:p")
                run = OxmlElement("w:r")
                text = OxmlElement("w:t")
                text.text = f"• {bullet}"  # Utilise une vraie puce ici
                run.append(text)
                new_para.append(run)
                parent.insert(index, new_para)

def modifie(doc, dict):
    for key in dict.keys():
        if isinstance(dict[key], list):
            texte_a_remplacer = f"[{key}]"
            liste_elements = dict[key]
            modifie_doc_bullet(word, texte_a_remplacer, liste_elements)
        else:
            remplace = key
            remplacant = dict[key]
            modifie_doc(doc, remplace, remplacant)

# donne le dictionnaire avec comme clé le nom du lot et comme valeur le prix (colonne 12)

def ajout_dico(df, colonne):
    list = []
    colonne_prix = df.iloc[:, colonne]
    i = 0
    while colonne_prix[i] != "Total arrondi":
        i = i+1
    i = i+1
    for j in range(i + 1, len(colonne_prix)):
        if isinstance(colonne_prix[j], int) and not isinstance(colonne_prix[j], bool):
            list.append(df_devis.iloc[j, 0] + f" : {colonne_prix[j]} euros")
    return list

Prix = ajout_dico(df_devis, 12)

def main():
    texte_a_remplacer = "[Axes_cadrage]"
    liste_elements = Prix
    modifie_doc_bullet(word, texte_a_remplacer, liste_elements)

    modifie(word, Contexte_Reference)
    modifie(word, Axes_remplis)
    Word_note(word)

    fichier_sortie = docs.sortie_path
    word.save(fichier_sortie)
    print("Document généré :", fichier_sortie)