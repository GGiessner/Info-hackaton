import re
from docx import Document
import docs

## Paramètres: Attention, il faut remplacer le nom des dossier par le nom des fichers de la mission

NOTES_FILE = docs.notes_3_path
WORD_TEMPLATE = "Template_mod.docx"  # Modèle Word
WORD_OUTPUT = docs.sortie_path
doc = Document(WORD_TEMPLATE)

## Infos notes
"""
notes_data = {}
with open(NOTES_FILE, 'r') as f:
    for line in f:
        if "::" in line:
            ident, content = line.strip().split("::", 1)
            notes_data[ident] = content

## code note

for paragraph in doc.paragraphs:
    for match in re.finditer(r'\{1(\w+)\}', paragraph.text):
        ident = match.group(1)
        if ident in notes_data:
            paragraph.text = paragraph.text.replace(match.group(0), notes_data[ident])

# Sauvegarder nouveau Word

doc.save(WORD_OUTPUT)
print(f" Nouveau document créé : {WORD_OUTPUT}")
"""

# Autre possibilité

from docx import Document

# Chemins des fichiers
NOTES_FILE = docs.notes_3_path
WORD_TEMPLATE = "Template.docx"
WORD_OUTPUT = docs.sortie_path

# Fonction pour lire les données du fichier texte
def lire_notes(fichier):
    notes_data = {}
    with open(fichier, 'r', encoding='utf-8') as f:
        for line in f:
            if "::" in line:
                ident, content = line.strip().split("::", 1)
                notes_data[ident.strip()] = content.strip()
    return notes_data

# Fonction pour modifier le document Word
def modifie_doc(doc, remplace, remplacant):
    for paragraph in doc.paragraphs:
        if remplace in paragraph.text:
            paragraph.text = paragraph.text.replace(remplace, remplacant)

# Fonction pour créer le contrat
def creation_contrat(doc, notes_data, doc_sortie):
    for ident, content in notes_data.items():
        modifie_doc(doc, f"{{1{ident}}}", content)
    doc.save(doc_sortie)

# Charger les données depuis le fichier texte
notes_data = lire_notes(NOTES_FILE)

# Charger le document Word
doc = Document(WORD_TEMPLATE)

# Créer le contrat
creation_contrat(doc, notes_data, WORD_OUTPUT)

print(f"Nouveau document créé : {WORD_OUTPUT}")
