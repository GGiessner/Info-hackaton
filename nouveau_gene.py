import json
import pandas as pd
import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.document import Document as DocumentClass

# Lecture Word (notes, propositions)
def lire_docx_en_texte(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

# Lecture Excel (devis)
def lire_excel_devis(path):
    df = pd.read_excel(path)
    devis_lignes = []
    for _, row in df.iterrows():
        devis_lignes.append([str(row[col]) if pd.notnull(row[col]) else "" for col in df.columns])
    return devis_lignes

def format_devis_text(devis_lignes):
    return "\n".join(["; ".join(ligne) for ligne in devis_lignes])

# Extraction structure document Word en liste (proposition exemple)
def extraire_structure_proposition(docx_path):
    doc = Document(docx_path)
    structure = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            style = block.style.name
            texte = block.text.strip()
            if not texte:
                continue
            if style.startswith("Heading"):
                structure.append({"type": "heading", "level": int(style[-1]), "text": texte})
            else:
                structure.append({"type": "paragraph", "text": texte})
        elif isinstance(block, Table):
            table_data = []
            for row in block.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            structure.append({"type": "table", "data": table_data})
    return structure

def iter_block_items(parent):
    if isinstance(parent, DocumentClass):
        parent_elm = parent.element.body
    else:
        raise ValueError("Parent doit être Document")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

# Prompt few-shot
def generer_prompt(notes_ex1, devis_ex1, prop_ex1_sections,
                   notes_ex2, devis_ex2, prop_ex2_sections,
                   notes_courant, devis_courant):
    prompt = f"""
Tu es un assistant expert qui rédige des propositions d'intervention pour une entreprise de conseil.

Voici deux exemples complets :

Exemple 1:
Notes:
{notes_ex1}

Devis:
{format_devis_text(devis_ex1)}

Proposition :
{json.dumps(prop_ex1_sections, ensure_ascii=False)}

Exemple 2:
Notes:
{notes_ex2}

Devis:
{format_devis_text(devis_ex2)}

Proposition :
{json.dumps(prop_ex2_sections, ensure_ascii=False)}

À partir des notes et devis suivants, génère une proposition structurée au format JSON avec des éléments :
- heading (avec level)
- paragraph (liste de textes)
- table (si nécessaire)

Notes :
{notes_courant}

Devis :
{format_devis_text(devis_courant)}
"""
    return prompt

# Appel à Ollama via Mistral
def generer_proposition(prompt):
    url = "http://localhost:11434/api/generate"
    payload = {
        "model": "mistral",
        "prompt": prompt,
        "stream": False
    }
    response = requests.post(url, json=payload)
    try:
        return response.json()['response']
    except Exception as e:
        print("Erreur lors du parsing JSON :", e)
        print("Texte brut :", response.text)
        raise

# Conversion format brut Mistral → format standard
def convertir_structure_mistral(donnees):
    if isinstance(donnees, str):
        donnees = json.loads(donnees)
    sections = donnees.get("Proposition d'Intervention", [])
    sortie = []
    for bloc in sections:
        if "heading" in bloc:
            sortie.append({
                "type": "heading",
                "level": bloc["heading"].get("level", 1),
                "text": bloc["heading"]["text"]
            })
        if "paragraph" in bloc:
            for p in bloc["paragraph"]:
                sortie.append({
                    "type": "paragraph",
                    "text": p["text"]
                })
        if "table" in bloc:
            sortie.append({
                "type": "table",
                "data": bloc["table"]
            })
    return sortie

# Construction du Word final
def construire_docx(structure, output_path):
    doc = Document()
    for elem in structure:
        type_elem = elem.get("type")
        if type_elem == "heading":
            doc.add_heading(elem["text"], level=elem.get("level", 1))
        elif type_elem == "paragraph":
            p = doc.add_paragraph(elem["text"])
            for run in p.runs:
                run.font.size = Pt(11)
        elif type_elem == "table":
            data = elem["data"]
            if not data:
                continue
            table = doc.add_table(rows=1, cols=len(data[0]))
            table.style = "Table Grid"
            for i, val in enumerate(data[0]):
                table.rows[0].cells[i].text = val
            for row in data[1:]:
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = val
    doc.save(output_path)

# === MAIN ===
if __name__ == "__main__":
    notes_ex1_path = "Gjoa MdP/Exemple 1/Exemple 1 - notes FC.docx"
    devis_ex1_path = "Gjoa MdP/Exemple 1/20250631-Client-Sujet-Devis exemple 1.xlsx"
    prop_ex1_path = "Gjoa MdP/Exemple 1/20250631-Gjoa-Sujet-Proposition.docx"

    notes_ex2_path = "Gjoa MdP/Exemple 2/Exemple 2 - notes FC.docx"
    devis_ex2_path = "Gjoa MdP/Exemple 2/20250401-Devis-Client-Exemple 2.xlsx"
    prop_ex2_path = "Gjoa MdP/Exemple 2/20250402-Gjoa-Client-Exemple 2-Proposition.docx"

    notes_3_path = "Gjoa MdP/Test/Test - notes FC.docx"
    devis_3_path = "Gjoa MdP/Test/20250515-Devis-client-projet.xlsx"

    sortie_path = "proposition_3_generée.docx"

    prop_ex1_sections = extraire_structure_proposition(prop_ex1_path)
    prop_ex2_sections = extraire_structure_proposition(prop_ex2_path)

    notes_ex1 = lire_docx_en_texte(notes_ex1_path)
    devis_ex1 = lire_excel_devis(devis_ex1_path)

    notes_ex2 = lire_docx_en_texte(notes_ex2_path)
    devis_ex2 = lire_excel_devis(devis_ex2_path)

    notes_3 = lire_docx_en_texte(notes_3_path)
    devis_3 = lire_excel_devis(devis_3_path)

    prompt = generer_prompt(notes_ex1, devis_ex1, prop_ex1_sections,
                            notes_ex2, devis_ex2, prop_ex2_sections,
                            notes_3, devis_3)

    print("Envoi du prompt à Mistral...")
    texte_recu = generer_proposition(prompt)
    print("Réponse reçue, génération du Word...")

    try:
        structure = convertir_structure_mistral(texte_recu)
    except Exception as e:
        print("Erreur lors de la conversion :", e)
        print("Texte reçu brut :", texte_recu)
        exit(1)

    construire_docx(structure, sortie_path)
    print(f"[OK] Document Word généré : {sortie_path}")
