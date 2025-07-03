import os
import docx
import json
import openpyxl
import requests
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd

# === FONCTIONS UTILITAIRES ===

def lire_docx(path):
    """Lit un fichier DOCX et retourne le texte brut concaténé."""
    doc = docx.Document(path)
    return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

def lire_xlsx(path):
    """Lit un fichier Excel et retourne le contenu concaténé des cellules non vides."""
    wb = openpyxl.load_workbook(path, data_only=True)
    contenu = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            ligne = [str(cell) for cell in row if cell is not None]
            if ligne:
                contenu.append(" | ".join(ligne))
    return "\n".join(contenu)

def extraire_texte_exemple(notes_path, devis_path, prop_path):
    notes = lire_docx(notes_path)
    devis = lire_xlsx(devis_path)
    prop = lire_docx(prop_path)
    return [notes, devis, prop]

def extraire_texte_test(notes_path, devis_path, template_path):
    notes = lire_docx(notes_path)
    devis = lire_xlsx(devis_path)
    template = lire_docx(template_path)
    return [notes, devis, template]

def generer_prompt(ex1, ex2, test_input, sections_demandées):
    sections = "\n".join(f"- {s}" for s in sections_demandées)
    return f"""
Tu es un assistant chargé de rédiger des propositions d'intervention professionnelles à partir de notes et de devis.

Voici deux exemples complets :

### Exemple 1 :
notes: {ex1[0]}
contrat: {ex1[2]}

### Exemple 2 :
notes: {ex2[0]}
contrat: {ex2[2]}

---

Voici maintenant l'étude de cas avec seulement les notes. C'est l'étude travaillée. C'EST CE SUR QUOI TU DOIS T'APPUYER POUR CE QUE JE TE DEMANDE PAR LA SUITE:
notes: {test_input[0]}

Génère uniquement les sections suivantes de la future proposition :
{sections}

Donne-moi la réponse au format JSON, avec un dictionnaire de la forme :
{{
  "Nom de la section 1": "Texte généré pour cette section",
  "Nom de la section 2": "..."
}}

Pour "Contexte":
Tu t'appuieras uniquement sur la note de l'étude travaillée.
Faire au moins 20 phrases.
Tu dois absolument inclure des données chiffrées, qui sont dans les notes
Indiquer la situation passée et actuelle du client.
Indiquer les besoins, les envies du client (expliqués en termes professionnels, dans le cadre d'une étude en conseil stratégique).
Je ne veux pas de réponse bateau. Je ne veux pas d'une réponse qui semble être applicable à toutes les entreprises, et qui est générale.
Je veux une réponse personnalisée!
"""

def interroger_mistral(prompt):
    """Envoie le prompt à Ollama/Mistral et renvoie la réponse décodée en dict."""
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": "mistral", "prompt": prompt, "stream": False}
        )
        response.raise_for_status()
        texte = response.json()["response"]
    except requests.RequestException as e:
        print(f"Erreur lors de la requête à Mistral : {e}")
        return {}

    try:
        return json.loads(texte)
    except json.JSONDecodeError:
        json_start = texte.find("{")
        json_end = texte.rfind("}") + 1
        try:
            return json.loads(texte[json_start:json_end])
        except:
            print("Erreur de parsing JSON. Contenu reçu :\n", texte)
            return {}

def enregistrer_resultats_word(resultats, sortie_path):
    doc = docx.Document()

    for section, texte in resultats.items():
        doc.add_heading(section, level=1)
        for paragraphe in texte.split("\n"):
            if paragraphe.strip():
                doc.add_paragraph(paragraphe.strip())
        doc.add_page_break()

    doc.save(sortie_path)
    print(f"Document Word généré : {sortie_path}")

def trouver_descriptions_proches(fichier_excel, texte_ref, top_n=9):
    colonnes_utiles = [
        "Description", 
        "Type de projet", 
        "Catégorie de projet", 
        "Mots clés (secteur, fonction, sujet)"
    ]
    df = pd.read_excel(fichier_excel)
    df = df[colonnes_utiles].dropna(subset=["Description"])

    corpus = [texte_ref] + df[colonnes_utiles].astype(str).agg(" ".join, axis=1).tolist()
    vect = TfidfVectorizer().fit_transform(corpus)
    sim_matrix = cosine_similarity(vect[0:1], vect[1:])
    scores = sim_matrix[0]
    top_indices = scores.argsort()[-top_n:][::-1]
    meilleures_descriptions = df.iloc[top_indices]["Description"].tolist()
    return meilleures_descriptions

def ajouter_descriptions_word(fichier_word, descriptions):
    doc = docx.Document(fichier_word)
    doc.add_page_break()
    doc.add_heading("Descriptions similaires issues du catalogue", level=1)
    for i, desc in enumerate(descriptions, 1):
        doc.add_paragraph(f"{i}. {desc}")
    doc.save(fichier_word)
    print(f"Descriptions ajoutées à {fichier_word}")

# === CHEMINS DES FICHIERS ===

notes_ex1_path = "Gjoa MdP/Exemple 1/Exemple 1 - notes FC.docx"
devis_ex1_path = "Gjoa MdP/Exemple 1/20250631-Client-Sujet-Devis exemple 1.xlsx"
prop_ex1_path  = "Gjoa MdP/Exemple 1/20250631-Gjoa-Sujet-Proposition.docx"

notes_ex2_path = "Gjoa MdP/Exemple 2/Exemple 2 - notes FC.docx"
devis_ex2_path = "Gjoa MdP/Exemple 2/20250401-Devis-Client-Exemple 2.xlsx"
prop_ex2_path  = "Gjoa MdP/Exemple 2/20250402-Gjoa-Client-Exemple 2-Proposition.docx"

notes_3_path   = "Gjoa MdP/Test/Test - notes FC.docx"
devis_3_path   = "Gjoa MdP/Test/20250515-Devis-client-projet.xlsx"
Template       = "Template.docx"
catalogue_path = "Gjoa MdP/202209-Liste références et bilans projets - extract MdP.xlsx"

# === SECTIONS À GÉNÉRER ===

sections_voulues = {
    "Contexte": None
}

# === PIPELINE PRINCIPAL ===

if __name__ == "__main__":
    print("Lecture des exemples...")
    exemple1 = extraire_texte_exemple(notes_ex1_path, devis_ex1_path, prop_ex1_path)
    exemple2 = extraire_texte_exemple(notes_ex2_path, devis_ex2_path, prop_ex2_path)
    test_input = extraire_texte_test(notes_3_path, devis_3_path, Template)

    print("Génération du prompt...")
    prompt = generer_prompt(exemple1, exemple2, test_input, sections_voulues.keys())

    print("Appel à Mistral via Ollama...")
    resultats = interroger_mistral(prompt)

    if "Contexte" in resultats:
        print("Recherche de descriptions similaires...")
        descriptions_proches = trouver_descriptions_proches(
            catalogue_path, texte_ref=resultats["Contexte"], top_n=9
        )
        ajouter_descriptions_word("proposition_test.docx", descriptions_proches)
