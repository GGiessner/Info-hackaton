import re
from docx import Document

## Paramètres: Attention, il faut remplacer le nom des dossier par le nom des fichers de la mission

NOTES_FILE = "C:/Users/guill/OneDrive/Documents/Scolaire/projet-info/Info-hackaton/Test notes/test_notes.txt"         # Notes brutes
WORD_TEMPLATE =  "C:/Users/guill/OneDrive/Documents/Scolaire/projet-info/Info-hackaton/Test notes/template_mod.docx"  # Modèle Word
WORD_OUTPUT =  "C:/Users/guill/OneDrive/Documents/Scolaire/projet-info/Info-hackaton/Test notes/result.docx"      # Fichier final
doc = Document(WORD_TEMPLATE)

## Infos notes

notes_data = {}
with open(NOTES_FILE, 'r', encoding='utf-8') as f:
    for line in f:
        if "::" in line:
            parts = line.strip().split("::", 1)
            if len(parts) == 2:
                ident, content = parts
                notes_data[ident.strip()] = content.strip()

print("Contenu de notes_data :", notes_data)

# Ouvrir le document Word
doc = Document(WORD_TEMPLATE)

# Remplacement des placeholders
for paragraph in doc.paragraphs:
    for match in re.finditer(r'\{1([^}]+)\}', paragraph.text):
        ident = match.group(1).strip()
        if ident in notes_data:
            paragraph.text = paragraph.text.replace(match.group(0), notes_data[ident])

# Sauvegarder le nouveau document Word
doc.save(WORD_OUTPUT)

print(f"Nouveau document créé : {WORD_OUTPUT}")