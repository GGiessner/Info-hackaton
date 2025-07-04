from docx import Document
import docs

# Chemins des fichiers
docs_path = docs.template_path
word_path = docs.test_notes_path
sortie_path = docs.sortie_path

def Word_note():
    # Charger le document Word contenant les associations
    associations_doc = Document(word_path)  # Utilisez word_path au lieu de associations_docx_path

    # Lire les associations identifiant-valeur depuis le document Word
    associations = {}
    for paragraph in associations_doc.paragraphs:
        if '::' in paragraph.text:
            parts = paragraph.text.split('::')
            if len(parts) == 2:  # Assurez-vous que la division a réussi
                identifiant, valeur = parts
                associations[identifiant.strip()] = valeur.strip()

    # Charger le document Word à modifier
    doc = Document(docs_path)

    # Fonction pour remplacer les identifiants dans le texte
    def replace_identifiers(text):
        for identifiant, valeur in associations.items():
            placeholder = "[" + identifiant + "]"
            if placeholder in text:
                text = text.replace(placeholder, valeur)
        return text

    # Parcourir chaque paragraphe du document
    for paragraph in doc.paragraphs:
        paragraph.text = replace_identifiers(paragraph.text)

    # Parcourir chaque section du document pour les en-têtes
    for section in doc.sections:
        for header_paragraph in section.header.paragraphs:
            header_paragraph.text = replace_identifiers(header_paragraph.text)

    # Parcourir chaque tableau du document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = replace_identifiers(paragraph.text)


Word_note()