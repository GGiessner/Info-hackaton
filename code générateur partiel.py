import pandas as pd
import re
from docx import Document
from openai import OpenAI

## Paramètres: Attention, il faut remplacer le nom des dossier par le nom des fichers de la mission

EXCEL_FILE = "data.xlsx"         # Tableau de données
NOTES_FILE = "notes.txt"         # Notes brutes
WORD_TEMPLATE = "template.docx"  # Modèle Word
WORD_OUTPUT = "result.docx"      # Fichier final
doc = Document(WORD_TEMPLATE)

## Infos Excel

df = pd.read_excel(EXCEL_FILE)
excel_data = dict(zip(df['identifiant'], df['valeur']))

## Infos notes

notes_data = {}
with open(NOTES_FILE, 'r') as f:
    for line in f:
        if "::" in line:
            ident, content = line.strip().split("::", 1)
            notes_data[ident] = content

## code note

for paragraph in doc.paragraphs:
    for match in re.finditer(r'\{1(\w+)\}', paragraph.text):
        ident = match.group(1)
        #if ident in excel_data:
        #    paragraph.text = paragraph.text.replace(match.group(0), excel_data[ident])
        if ident in notes_data:
            paragraph.text = paragraph.text.replace(match.group(0), notes_data[ident])



## code excel



# Préparation IA

client = OpenAI()

## code IA




# Remplacer les identifiants 

def replace_in_paragraph(paragraph):
    matches = re.findall(r"\[(.+?)\]", paragraph.text)
    for ident in matches:
        if ident in excel_data:
            paragraph.text = paragraph.text.replace(f"[{ident}]", str(excel_data[ident]))
        elif ident in notes_data:
            if ident.startswith("IA_"):
                prompt = notes_data[ident]
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": "Tu rédiges des paragraphes clairs et structurés."},
                        {"role": "user", "content": f"Voici les notes : {prompt}. Rédige un paragraphe rédigé formel."}
                    ]
                )
                generated = response.choices[0].message.content
                paragraph.text = paragraph.text.replace(f"[{ident}]", generated)
            else:
                paragraph.text = paragraph.text.replace(f"[{ident}]", notes_data[ident])

# Parcourir tout le document

for paragraph in doc.paragraphs:
    replace_in_paragraph(paragraph)

# Sauvegarder nouveau Word

doc.save(WORD_OUTPUT)


print(f" Nouveau document créé : {WORD_OUTPUT}")

